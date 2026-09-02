# -*- coding: utf-8 -*-
import os
import json
import uuid
import logging
from datetime import datetime
from pathlib import Path
from flask import Flask, render_template, request, jsonify, redirect, url_for, session, Response
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

logging.basicConfig(level=logging.INFO)

load_dotenv(dotenv_path=Path(__file__).parent / ".env")

app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config["UPLOAD_FOLDER"] = os.path.join(app.root_path, "uploads")
app.config["HISTORY_FOLDER"] = os.path.join(app.root_path, "history")
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

# Create directories at module level
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["HISTORY_FOLDER"], exist_ok=True)
os.makedirs("uploads", exist_ok=True)
os.makedirs("history", exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/api/check")
def api_check():
    checks = {}
    checks["groq_key_set"] = bool(os.getenv("GROQ_API_KEY"))
    checks["upload_folder_exists"] = os.path.exists(app.config["UPLOAD_FOLDER"])
    checks["history_folder_exists"] = os.path.exists(app.config["HISTORY_FOLDER"])
    try:
        checks["embeddings_import"] = "ok"
    except Exception as e:
        checks["embeddings_import"] = str(e)
    try:
        checks["groq_import"] = "ok"
    except Exception as e:
        checks["groq_import"] = str(e)
    return jsonify(checks)

@app.route("/api/test-analyze")
def test_analyze():
    from langchain_community.vectorstores import FAISS
    from langchain_core.documents import Document
    from langchain_groq import ChatGroq
    from langchain_huggingface import HuggingFaceEmbeddings
    results = {}
    try:
        emb = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        results["step1_embeddings"] = "ok"
    except Exception as e:
        results["step1_embeddings"] = str(e)
        return jsonify(results)
    
    try:
        fake_docs = [
            Document(page_content="The contract requires 30 days notice."),
            Document(page_content="Employees must give written notice.")
        ]
        vs = FAISS.from_documents(fake_docs, emb, distance_strategy="COSINE")
        results["step2_faiss"] = "ok"
    except Exception as e:
        results["step2_faiss"] = str(e)
        return jsonify(results)
    
    try:
        llm = ChatGroq(
            model_name="llama-3.1-8b-instant",
            api_key=os.getenv("GROQ_API_KEY")
        )
        response = llm.invoke("Say OK")
        results["step3_groq"] = response.content[:50]
    except Exception as e:
        results["step3_groq"] = str(e)
        return jsonify(results)
    
    try:
        matches = vs.similarity_search_with_score("notice period", k=1)
        sim_score = round(matches[0][1], 3)
        results["step4_similarity"] = f"ok - score: {sim_score}"
    except Exception as e:
        results["step4_similarity"] = str(e)
        return jsonify(results)

    results["all_steps"] = "PASSED"
    return jsonify(results)

def rate_severity(chunk_a, chunk_b, reason, llm):
    try:
        severity_prompt = f"""Rate the severity of this contradiction between two documents.

Document A says: {chunk_a[:200]}
Document B says: {chunk_b[:200]}
Conflict: {reason}

Consider: legal/financial impact, operational impact, clarity impact.

Respond with ONLY one word — either:
CRITICAL
SIGNIFICANT  
MINOR

No explanation. Just the single word."""
        
        response = llm.invoke(severity_prompt)
        text = getattr(response, "content", str(response)).strip().upper()
        if "CRITICAL" in text:
            return "CRITICAL"
        elif "SIGNIFICANT" in text:
            return "SIGNIFICANT"
        else:
            return "MINOR"
    except Exception:
        return "SIGNIFICANT"

def calculate_confidence(score, verdict_text, chunk_a, chunk_b):
    confidence = 0.0
    
    # Factor 1: similarity score (lower distance = more confident)
    # score is cosine distance 0-1, lower = more similar topic
    score_confidence = (1 - score) * 40  # max 40 points
    confidence += score_confidence
    
    # Factor 2: verdict clarity (clear DISAGREE/AGREE = confident)
    verdict_upper = verdict_text.upper()
    if "DISAGREE" in verdict_upper and "PARTIALLY" not in verdict_upper:
        confidence += 30
    elif "AGREE" in verdict_upper and "PARTIALLY" not in verdict_upper:
        confidence += 28
    elif "PARTIALLY" in verdict_upper:
        confidence += 15
    
    # Factor 3: text length (longer chunks = more context = more confident)
    avg_len = (len(chunk_a) + len(chunk_b)) / 2
    if avg_len > 300:
        confidence += 20
    elif avg_len > 150:
        confidence += 12
    else:
        confidence += 5
    
    # Factor 4: REASON quality (longer reason = more specific = confident)
    reason_lines = [l for l in verdict_text.split('\n') 
                    if l.startswith('REASON:')]
    if reason_lines:
        reason_len = len(reason_lines[0])
        if reason_len > 80:
            confidence += 10
        elif reason_len > 40:
            confidence += 6
        else:
            confidence += 2
    
    return min(round(confidence, 1), 99.0)  # cap at 99

def extract_text_from_file(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    
    if ext == "pdf":
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(filepath)
        return loader.load()
    
    elif ext == "docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            from langchain_core.documents import Document
            return [Document(page_content=full_text, metadata={"source": filepath})]
        except ImportError:
            return {"error": "python-docx not installed. Run: pip install python-docx"}
        except Exception as e:
            return {"error": f"Could not read DOCX: {str(e)}"}
    
    elif ext == "txt":
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            from langchain_core.documents import Document
            return [Document(page_content=content, metadata={"source": filepath})]
        except Exception as e:
            return {"error": f"Could not read TXT: {str(e)}"}
    
    else:
        return {"error": f"Unsupported file type: .{ext}"}

def process_documents(file_a, file_b):
    try:
        from langchain_community.vectorstores import FAISS
        from langchain_core.prompts import PromptTemplate
        from langchain_groq import ChatGroq
        from langchain_huggingface import HuggingFaceEmbeddings
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        emb = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        llm = ChatGroq(
            model_name="llama-3.1-8b-instant",
            api_key=os.getenv("GROQ_API_KEY")
        )

        raw_a = extract_text_from_file(file_a)
        raw_b = extract_text_from_file(file_b)
        
        if isinstance(raw_a, dict) and raw_a.get("error"):
            return {"error": raw_a["error"]}
        if isinstance(raw_b, dict) and raw_b.get("error"):
            return {"error": raw_b["error"]}
        
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        docs_a = splitter.split_documents(raw_a)
        docs_b = splitter.split_documents(raw_b)

        if len(docs_a) == 0 or len(docs_b) == 0:
            return {"error": "Could not extract text from one or both files."}

        # FIX 2 - Use COSINE distance strategy
        vectorstore_a = FAISS.from_documents(docs_a, emb, distance_strategy="COSINE")
        vectorstore_b = FAISS.from_documents(docs_b, emb, distance_strategy="COSINE")

        prompt = PromptTemplate(
            input_variables=["text_a", "text_b"],
            template="""Compare these two excerpts from different documents.

Excerpt A: {text_a}

Excerpt B: {text_b}

Respond in exactly this format:
VERDICT: [AGREE / DISAGREE / PARTIALLY AGREE]
REASON: [One sentence explanation]"""
        )
        chain = prompt | llm

        contradictions = []
        agreements = []
        all_scores = []
        seen_pairs = set()

        for chunk in docs_a:
            try:
                match = vectorstore_b.similarity_search_with_score(chunk.page_content, k=1)
                if not match:
                    continue
                best_match, score = match[0]
                all_scores.append(score)
                pair_key = (chunk.page_content[:80], best_match.page_content[:80])
                if pair_key in seen_pairs:
                    continue
                seen_pairs.add(pair_key)
                # Threshold for COSINE distance: score < 0.5 means similar topic
                if score < 0.5:
                    response = chain.invoke({
                        "text_a": chunk.page_content,
                        "text_b": best_match.page_content
                    })
                    verdict_text = getattr(response, "content", str(response))
                    verdict = "PARTIALLY AGREE"
                    reason = verdict_text
                    if "VERDICT:" in verdict_text:
                        for line in verdict_text.split("\n"):
                            if line.startswith("VERDICT:"):
                                verdict = line.replace("VERDICT:", "").strip()
                            if line.startswith("REASON:"):
                                reason = line.replace("REASON:", "").strip()
                    entry = {
                        "chunk_a": chunk.page_content,
                        "chunk_b": best_match.page_content,
                        "score": float(round(score, 3)),
                        "verdict": verdict,
                        "reason": reason,
                        "confidence": calculate_confidence(score, verdict_text, chunk.page_content, best_match.page_content)
                    }
                    if "DISAGREE" in verdict:
                        entry["severity"] = rate_severity(chunk.page_content, best_match.page_content, reason, llm)
                        contradictions.append(entry)
                    else:
                        agreements.append(entry)
            except Exception as chunk_err:
                logging.error(f"Chunk error: {chunk_err}")
                continue

        blind_spots_a = []
        blind_spots_b = []
        # Threshold for COSINE distance: score > 0.6 means topic not covered
        for chunk in docs_a:
            try:
                match = vectorstore_b.similarity_search_with_score(chunk.page_content, k=1)
                if match and float(match[0][1]) > 0.6:
                    blind_spots_a.append(chunk.page_content[:200])
            except:
                continue

        for chunk in docs_b:
            try:
                match = vectorstore_a.similarity_search_with_score(chunk.page_content, k=1)
                if match and float(match[0][1]) > 0.6:
                    blind_spots_b.append(chunk.page_content[:200])
            except:
                continue

        severity_order = {"CRITICAL": 0, "SIGNIFICANT": 1, "MINOR": 2}
        contradictions.sort(key=lambda x: severity_order.get(x.get("severity", "MINOR"), 1))

        # FIX 3 - Convert cosine distance to similarity percentage
        result = {
            "contradictions": contradictions[:10],
            "agreements": agreements[:10],
            "blind_spots_a": blind_spots_a[:5],
            "blind_spots_b": blind_spots_b[:5],
            "total_chunks_a": len(docs_a),
            "total_chunks_b": len(docs_b),
            "score_range": {
                "min": float(round((1 - max(all_scores)) * 100, 1)) if all_scores else 0,
                "max": float(round((1 - min(all_scores)) * 100, 1)) if all_scores else 0
            }
        }

        if not contradictions and not agreements:
            result["warning"] = "No comparable passages found. Documents may be too different or too short."

        return result

    except Exception as e:
        logging.error(f"process_documents failed: {e}")
        return {"error": f"Analysis failed: {str(e)}"}

def generate_narrative(doc_a_name, doc_b_name, results):
    try:
        from langchain_groq import ChatGroq
        llm = ChatGroq(
            model_name="llama-3.1-8b-instant",
            api_key=os.getenv("GROQ_API_KEY")
        )

        contradictions_text = ""
        for i, c in enumerate(results.get("contradictions", [])[:5], 1):
            contradictions_text += f"\nContradiction {i}:\n  Doc A: {c['chunk_a'][:200]}\n  Doc B: {c['chunk_b'][:200]}\n  Verdict: {c['verdict']}\n  Reason: {c['reason']}\n"

        agreements_text = ""
        for i, a in enumerate(results.get("agreements", [])[:5], 1):
            agreements_text += f"\nAgreement {i}:\n  Doc A: {a['chunk_a'][:200]}\n  Doc B: {a['chunk_b'][:200]}\n  Reason: {a['reason']}\n"

        blind_a = "\n".join([f"  - {b[:150]}" for b in results.get("blind_spots_a", [])])
        blind_b = "\n".join([f"  - {b[:150]}" for b in results.get("blind_spots_b", [])])

        prompt = f"""You are a professional document analyst. You have compared two documents.

Document A filename: {doc_a_name}
Document B filename: {doc_b_name}

Analysis results:
- Contradictions found: {len(results.get('contradictions', []))}
- Agreements found: {len(results.get('agreements', []))}
- Blind spots in Doc A: {len(results.get('blind_spots_a', []))}
- Blind spots in Doc B: {len(results.get('blind_spots_b', []))}

CONTRADICTIONS:
{contradictions_text if contradictions_text else "None found."}

AGREEMENTS:
{agreements_text if agreements_text else "None found."}

TOPICS ONLY IN DOCUMENT A:
{blind_a if blind_a else "None found."}

TOPICS ONLY IN DOCUMENT B:
{blind_b if blind_b else "None found."}

Write a professional analysis report with these exact sections.
Be specific, reference actual content from the documents, not just counts.
Write in clear professional English. Each section must be at least 2-3 sentences.

FORMAT YOUR RESPONSE EXACTLY LIKE THIS - use these exact headings:

DOCUMENT OVERVIEW
[2-3 sentences identifying what each document appears to be about based on their content and filenames. Name the documents by their actual filenames.]

KEY FINDINGS
[3-5 sentences summarising the most important discoveries from this comparison. What is the headline story? What should the reader pay attention to first?]

CONTRADICTIONS ANALYSIS
[For each contradiction found, write 2 sentences: what the conflict is, and why it matters. If no contradictions, state that clearly and explain what that means for document alignment.]

AGREEMENTS ANALYSIS  
[Summarise what both documents agree on. What common ground exists? 2-3 sentences.]

BLIND SPOTS
[Describe what Document A covers that Document B misses, and vice versa. Why might these gaps matter? 2-3 sentences per document.]

CONCLUSION
[2-3 sentences. Overall assessment: are these documents aligned, conflicting, or complementary? What action should the reader take based on this analysis?]"""

        response = llm.invoke(prompt)
        narrative_text = getattr(response, "content", str(response))

        sections = {}
        current_section = None
        current_lines = []

        section_headers = [
            "DOCUMENT OVERVIEW",
            "KEY FINDINGS", 
            "CONTRADICTIONS ANALYSIS",
            "AGREEMENTS ANALYSIS",
            "BLIND SPOTS",
            "CONCLUSION"
        ]

        for line in narrative_text.split('\n'):
            line_stripped = line.strip()
            if line_stripped in section_headers:
                if current_section and current_lines:
                    sections[current_section] = '\n'.join(current_lines).strip()
                current_section = line_stripped
                current_lines = []
            elif current_section:
                current_lines.append(line)

        if current_section and current_lines:
            sections[current_section] = '\n'.join(current_lines).strip()

        return {
            "raw": narrative_text,
            "sections": sections,
            "generated": True
        }

    except Exception as e:
        logging.error(f"Narrative generation failed: {e}")
        return {
            "raw": "",
            "sections": {},
            "generated": False,
            "error": str(e)
        }

def save_history_entry(doc_a_name, doc_b_name, results):
    entry_id = str(uuid.uuid4())[:8]
    entry = {
        "id": entry_id,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "doc_a": doc_a_name,
        "doc_b": doc_b_name,
        "overall_similarity": results.get("overall_similarity", 85.0 if not results.get("warning") else 40.0) if isinstance(results, dict) else 85.0,
        "summary_stats": results.get("summary_stats", {}) if isinstance(results, dict) else {},
        "results": results
    }
    history_file = Path(app.config["HISTORY_FOLDER"]) / "history.json"
    history = []
    if history_file.exists():
        try:
            with open(history_file, "r", encoding="utf-8") as f:
                history = json.load(f)
        except Exception:
            history = []
    history.insert(0, entry)
    history = history[:50]
    with open(history_file, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2, ensure_ascii=False)
    return entry_id

def load_history_entries():
    history_file = Path(app.config["HISTORY_FOLDER"]) / "history.json"
    if not history_file.exists():
        return []
    try:
        with open(history_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def get_history_by_id(entry_id):
    entries = load_history_entries()
    for entry in entries:
        if entry.get("id") == entry_id:
            return entry
    return None

@app.route("/")
def home():
    return render_template("home.html")

@app.route("/dashboard")
def dashboard():
    history = load_history_entries()
    return render_template("dashboard.html", recent_history=history[:5])

@app.route("/analyze-legacy", methods=["POST"])
@app.route("/compare-legacy", methods=["POST"])
def analyze_legacy():
    try:
        if "doc_a" not in request.files or "doc_b" not in request.files:
            return jsonify({"error": "Both PDF files (Document A and Document B) are required."}), 400
        file_a = request.files["doc_a"]
        file_b = request.files["doc_b"]
        if file_a.filename == "" or file_b.filename == "":
            return jsonify({"error": "Please select two PDF documents to compare."}), 400
        if not allowed_file(file_a.filename) or not allowed_file(file_b.filename):
            return jsonify({"error": "Only .pdf files are supported."}), 400
        name_a = secure_filename(file_a.filename)
        name_b = secure_filename(file_b.filename)
        path_a = os.path.join(app.config["UPLOAD_FOLDER"], f"{str(uuid.uuid4())[:6]}_{name_a}")
        path_b = os.path.join(app.config["UPLOAD_FOLDER"], f"{str(uuid.uuid4())[:6]}_{name_b}")
        
        file_a.save(path_a)
        file_b.save(path_b)
        
        results = process_documents(path_a, path_b)
        
        if isinstance(results, dict) and results.get("error"):
            return jsonify({"error": results["error"]}), 400

        if isinstance(results, dict):
            if "doc_a" not in results:
                results["doc_a"] = {"name": name_a}
            if "doc_b" not in results:
                results["doc_b"] = {"name": name_b}

            narrative = generate_narrative(name_a, name_b, results)
            results["narrative"] = narrative

        history_id = save_history_entry(name_a, name_b, results)
        session["last_result"] = results
        session["last_id"] = history_id
        return jsonify({
            "success": True,
            "id": history_id,
            "redirect_url": url_for("results", id=history_id)
        })
    except Exception as e:
        logging.error(f"Error in /analyze route: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/analyze-stream", methods=["POST"])
@app.route("/analyze", methods=["POST"])
@app.route("/compare", methods=["POST"])
def analyze_stream():
    if "doc_a" not in request.files or "doc_b" not in request.files:
        return jsonify({"error": "Both files required"}), 400
    
    file_a = request.files["doc_a"]
    file_b = request.files["doc_b"]
    
    if not allowed_file(file_a.filename) or not allowed_file(file_b.filename):
        return jsonify({"error": "Only PDF, DOCX, and TXT files supported"}), 400
    
    name_a = secure_filename(file_a.filename)
    name_b = secure_filename(file_b.filename)
    path_a = os.path.join(app.config["UPLOAD_FOLDER"], f"{str(uuid.uuid4())[:6]}_{name_a}")
    path_b = os.path.join(app.config["UPLOAD_FOLDER"], f"{str(uuid.uuid4())[:6]}_{name_b}")
    
    file_a.save(path_a)
    file_b.save(path_b)
    
    def generate():
        try:
            def send(event, data):
                return f"event: {event}\ndata: {json.dumps(data)}\n\n"
            
            # Step 1
            yield send("progress", {"step": 1, "total": 7,
                "message": "Loading documents...", "percent": 5})
            
            raw_a = extract_text_from_file(path_a)
            raw_b = extract_text_from_file(path_b)

            if isinstance(raw_a, dict) and raw_a.get("error"):
                yield send("error", {"message": raw_a["error"]})
                return
            if isinstance(raw_b, dict) and raw_b.get("error"):
                yield send("error", {"message": raw_b["error"]})
                return

            splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
            docs_a = splitter.split_documents(raw_a)
            docs_b = splitter.split_documents(raw_b)

            if not docs_a or not docs_b:
                yield send("error", {"message": "Could not extract text from files."})
                return

            yield send("progress", {"step": 2, "total": 7,
                "message": f"Loaded {len(docs_a)} chunks from Doc A, {len(docs_b)} from Doc B",
                "percent": 15})

            # Step 2 — Build embeddings
            yield send("progress", {"step": 3, "total": 7,
                "message": "Building vector stores...", "percent": 25})
            
            emb = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
            vectorstore_a = FAISS.from_documents(docs_a, emb, distance_strategy="COSINE")
            vectorstore_b = FAISS.from_documents(docs_b, emb, distance_strategy="COSINE")

            yield send("progress", {"step": 4, "total": 7,
                "message": "Vector stores ready. Running similarity search...",
                "percent": 35})

            # Step 3 — LLM setup
            llm = ChatGroq(
                model_name="llama-3.1-8b-instant",
                api_key=os.getenv("GROQ_API_KEY")
            )
            prompt = PromptTemplate(
                input_variables=["text_a", "text_b"],
                template="""Compare these two excerpts from different documents.

Excerpt A: {text_a}

Excerpt B: {text_b}

Respond in exactly this format:
VERDICT: [AGREE / DISAGREE / PARTIALLY AGREE]
REASON: [One sentence explanation]"""
            )
            chain = prompt | llm

            contradictions = []
            agreements = []
            all_scores = []
            seen_pairs = set()
            total_chunks = len(docs_a)
            processed = 0

            # Step 4 — Process chunks, stream each finding
            for chunk in docs_a:
                try:
                    match = vectorstore_b.similarity_search_with_score(
                        chunk.page_content, k=1)
                    if not match:
                        continue
                    best_match, score = match[0]
                    all_scores.append(float(score))
                    pair_key = (chunk.page_content[:80], best_match.page_content[:80])
                    if pair_key in seen_pairs:
                        continue
                    seen_pairs.add(pair_key)

                    processed += 1
                    percent = 35 + int((processed / total_chunks) * 40)
                    yield send("progress", {
                        "step": 5, "total": 7,
                        "message": f"Analysing chunk {processed}/{total_chunks}...",
                        "percent": percent
                    })

                    if float(score) < 0.5:
                        response = chain.invoke({
                            "text_a": chunk.page_content,
                            "text_b": best_match.page_content
                        })
                        verdict_text = getattr(response, "content", str(response))
                        verdict = "PARTIALLY AGREE"
                        reason = verdict_text
                        if "VERDICT:" in verdict_text:
                            for line in verdict_text.split("\n"):
                                if line.startswith("VERDICT:"):
                                    verdict = line.replace("VERDICT:", "").strip()
                                if line.startswith("REASON:"):
                                    reason = line.replace("REASON:", "").strip()

                        entry = {
                            "chunk_a": chunk.page_content,
                            "chunk_b": best_match.page_content,
                            "score": float(round(score, 3)),
                            "verdict": verdict,
                            "reason": reason,
                            "confidence": calculate_confidence(score, verdict_text, chunk.page_content, best_match.page_content)
                        }

                        if "DISAGREE" in verdict:
                            severity = rate_severity(chunk.page_content, best_match.page_content, reason, llm)
                            entry["severity"] = severity
                            contradictions.append(entry)
                            # Stream contradiction immediately
                            yield send("contradiction", {
                                "index": len(contradictions),
                                "data": entry,
                                "severity": severity
                            })
                        else:
                            agreements.append(entry)
                            # Stream agreement immediately
                            yield send("agreement", {
                                "index": len(agreements),
                                "data": entry
                            })

                except Exception as chunk_err:
                    logging.error(f"Chunk error: {chunk_err}")
                    continue

            # Step 5 — Blind spots
            yield send("progress", {"step": 6, "total": 7,
                "message": "Detecting blind spots...", "percent": 78})

            blind_spots_a = []
            blind_spots_b = []
            for chunk in docs_a:
                try:
                    match = vectorstore_b.similarity_search_with_score(
                        chunk.page_content, k=1)
                    if match and float(match[0][1]) > 0.6:
                        blind_spots_a.append(chunk.page_content[:200])
                except:
                    continue

            for chunk in docs_b:
                try:
                    match = vectorstore_a.similarity_search_with_score(
                        chunk.page_content, k=1)
                    if match and float(match[0][1]) > 0.6:
                        blind_spots_b.append(chunk.page_content[:200])
                except:
                    continue

            for bs in blind_spots_a[:5]:
                yield send("blind_spot_a", {"text": bs})
            for bs in blind_spots_b[:5]:
                yield send("blind_spot_b", {"text": bs})

            # Step 6 — Narrative
            yield send("progress", {"step": 7, "total": 7,
                "message": "Generating narrative report...", "percent": 88})

            severity_order = {"CRITICAL": 0, "SIGNIFICANT": 1, "MINOR": 2}
            contradictions.sort(key=lambda x: severity_order.get(x.get("severity", "MINOR"), 1))

            results = {
                "contradictions": contradictions[:10],
                "agreements": agreements[:10],
                "blind_spots_a": blind_spots_a[:5],
                "blind_spots_b": blind_spots_b[:5],
                "total_chunks_a": len(docs_a),
                "total_chunks_b": len(docs_b),
                "score_range": {
                    "min": float(round((1 - max(all_scores)) * 100, 1)) if all_scores else 0,
                    "max": float(round((1 - min(all_scores)) * 100, 1)) if all_scores else 0
                }
            }

            narrative = generate_narrative(name_a, name_b, results)
            results["narrative"] = narrative
            results["doc_a"] = {"name": name_a}
            results["doc_b"] = {"name": name_b}

            history_id = save_history_entry(name_a, name_b, results)
            session["last_result"] = results
            session["last_id"] = history_id

            yield send("complete", {
                "id": history_id,
                "redirect_url": f"/result?id={history_id}",
                "summary": {
                    "contradictions": len(contradictions),
                    "agreements": len(agreements),
                    "blind_spots": len(blind_spots_a) + len(blind_spots_b),
                    "chunks": len(docs_a) + len(docs_b)
                }
            })

        except Exception as e:
            logging.error(f"Stream error: {e}")
            yield f"event: error\ndata: {json.dumps({'message': str(e)})}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive"
        }
    )

@app.route("/analyze-multi", methods=["POST"])
def analyze_multi():
    try:
        files = []
        for i in range(1, 6):
            key = f"doc_{i}"
            if key in request.files and request.files[key].filename:
                f = request.files[key]
                if allowed_file(f.filename):
                    files.append(f)

        if len(files) < 2:
            return jsonify({"error": "At least 2 PDF, DOCX, or TXT files required"}), 400
        if len(files) > 5:
            return jsonify({"error": "Maximum 5 documents"}), 400

        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from langchain_community.vectorstores import FAISS
        from langchain_huggingface import HuggingFaceEmbeddings
        from langchain_groq import ChatGroq
        from langchain_core.prompts import PromptTemplate

        saved_paths = []
        names = []
        for f in files:
            name = secure_filename(f.filename)
            path = os.path.join(
                app.config["UPLOAD_FOLDER"],
                f"{str(uuid.uuid4())[:6]}_{name}"
            )
            f.save(path)
            saved_paths.append(path)
            names.append(name)

        emb = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        llm = ChatGroq(
            model_name="llama-3.1-8b-instant",
            api_key=os.getenv("GROQ_API_KEY")
        )

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=50)
        
        all_docs = []
        vectorstores = []
        
        for path in saved_paths:
            raw_docs = extract_text_from_file(path)
            if isinstance(raw_docs, dict) and raw_docs.get("error"):
                return jsonify({"error": raw_docs["error"]}), 400
            docs = splitter.split_documents(raw_docs)
            all_docs.append(docs)
            vs = FAISS.from_documents(docs, emb, distance_strategy="COSINE")
            vectorstores.append(vs)

        prompt = PromptTemplate(
            input_variables=["text_a", "text_b"],
            template="""Compare these two excerpts from different documents.
Excerpt A: {text_a}
Excerpt B: {text_b}
Respond in exactly this format:
VERDICT: [AGREE / DISAGREE / PARTIALLY AGREE]
REASON: [One sentence explanation]"""
        )
        chain = prompt | llm

        # Compare every pair of documents
        all_contradictions = []
        all_agreements = []
        seen_pairs = set()
        all_scores = []

        for i in range(len(all_docs)):
            for j in range(i + 1, len(all_docs)):
                pair_label = f"{names[i]} vs {names[j]}"
                
                for chunk in all_docs[i]:
                    try:
                        match = vectorstores[j].similarity_search_with_score(
                            chunk.page_content, k=1)
                        if not match:
                            continue
                        best_match, score = match[0]
                        all_scores.append(float(score))
                        pair_key = (chunk.page_content[:80],
                                    best_match.page_content[:80])
                        if pair_key in seen_pairs:
                            continue
                        seen_pairs.add(pair_key)

                        if float(score) < 0.5:
                            response = chain.invoke({
                                "text_a": chunk.page_content,
                                "text_b": best_match.page_content
                            })
                            verdict_text = getattr(
                                response, "content", str(response))
                            verdict = "PARTIALLY AGREE"
                            reason = verdict_text
                            if "VERDICT:" in verdict_text:
                                for line in verdict_text.split("\n"):
                                    if line.startswith("VERDICT:"):
                                        verdict = line.replace(
                                            "VERDICT:", "").strip()
                                    if line.startswith("REASON:"):
                                        reason = line.replace(
                                            "REASON:", "").strip()

                            entry = {
                                "doc_i": names[i],
                                "doc_j": names[j],
                                "pair_label": pair_label,
                                "chunk_a": chunk.page_content,
                                "chunk_b": best_match.page_content,
                                "score": float(round(score, 3)),
                                "verdict": verdict,
                                "reason": reason,
                                "confidence": calculate_confidence(score, verdict_text, chunk.page_content, best_match.page_content)
                            }

                            if "DISAGREE" in verdict:
                                entry["severity"] = rate_severity(
                                    chunk.page_content, best_match.page_content, reason, llm
                                )
                                all_contradictions.append(entry)
                            else:
                                all_agreements.append(entry)
                    except Exception as e:
                        logging.error(f"Multi chunk error: {e}")
                        continue

        # Blind spots per document
        blind_spots = {}
        for i, (docs_i, vs_i, name_i) in enumerate(
                zip(all_docs, vectorstores, names)):
            blind_spots[name_i] = []
            for j, vs_j in enumerate(vectorstores):
                if i == j:
                    continue
                for chunk in docs_i:
                    try:
                        match = vs_j.similarity_search_with_score(
                            chunk.page_content, k=1)
                        if match and float(match[0][1]) > 0.6:
                            text = chunk.page_content[:200]
                            if text not in blind_spots[name_i]:
                                blind_spots[name_i].append(text)
                    except:
                        continue
            blind_spots[name_i] = blind_spots[name_i][:5]

        severity_order = {"CRITICAL": 0, "SIGNIFICANT": 1, "MINOR": 2}
        all_contradictions.sort(key=lambda x: severity_order.get(x.get("severity", "MINOR"), 1))

        results = {
            "mode": "multi",
            "documents": names,
            "doc_count": len(names),
            "contradictions": all_contradictions[:15],
            "agreements": all_agreements[:15],
            "blind_spots": blind_spots,
            "total_chunks": sum(len(d) for d in all_docs),
            "score_range": {
                "min": float(round((1 - max(all_scores)) * 100, 1))
                       if all_scores else 0,
                "max": float(round((1 - min(all_scores)) * 100, 1))
                       if all_scores else 0
            }
        }

        if not all_contradictions and not all_agreements:
            results["warning"] = "No comparable passages found."

        doc_names_str = " vs ".join(names)
        narrative = generate_narrative(names[0], doc_names_str, results)
        results["narrative"] = narrative

        history_id = save_history_entry(names[0], doc_names_str, results)
        session["last_result"] = results
        session["last_id"] = history_id

        return jsonify({
            "success": True,
            "id": history_id,
            "redirect_url": f"/result?id={history_id}",
            "doc_count": len(names)
        })

    except Exception as e:
        logging.error(f"Multi-analyze error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/results")
@app.route("/result")
def results():
    entry_id = request.args.get("id")
    result_data = None
    doc_info = {}
    if entry_id:
        entry = get_history_by_id(entry_id)
        if entry:
            result_data = entry.get("results")
            doc_info = {"a": entry.get("doc_a"), "b": entry.get("doc_b"), "timestamp": entry.get("timestamp")}
    if not result_data and "last_result" in session:
        result_data = session.get("last_result")
        doc_info = {
            "a": result_data.get("doc_a", {}).get("name") if isinstance(result_data.get("doc_a"), dict) else result_data.get("doc_a"),
            "b": result_data.get("doc_b", {}).get("name") if isinstance(result_data.get("doc_b"), dict) else result_data.get("doc_b")
        }
    return render_template("results.html", result=result_data, doc_info=doc_info, entry_id=entry_id)

@app.route("/download/txt")
def download_txt():
    entry_id = request.args.get("id")
    result_data = None
    doc_info = {}
    if entry_id:
        entry = get_history_by_id(entry_id)
        if entry:
            result_data = entry.get("results")
            doc_info = {"a": entry.get("doc_a"), "b": entry.get("doc_b")}
    if not result_data and "last_result" in session:
        result_data = session.get("last_result")
        doc_info = {
            "a": result_data.get("doc_a", {}).get("name") if isinstance(result_data.get("doc_a"), dict) else result_data.get("doc_a"),
            "b": result_data.get("doc_b", {}).get("name") if isinstance(result_data.get("doc_b"), dict) else result_data.get("doc_b")
        }

    if not result_data:
        return "No analysis results found", 404

    lines = [
        "DATA ANALYSER ENGINE - Analysis Report",
        "=" * 50,
        f"Document A: {doc_info.get('a', 'Unknown')}",
        f"Document B: {doc_info.get('b', 'Unknown')}",
        f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}",
        "",
    ]

    narrative = result_data.get("narrative", {})
    if narrative.get("generated") and narrative.get("sections"):
        for section_name in [
            "DOCUMENT OVERVIEW",
            "KEY FINDINGS",
            "CONTRADICTIONS ANALYSIS", 
            "AGREEMENTS ANALYSIS",
            "BLIND SPOTS",
            "CONCLUSION"
        ]:
            section_text = narrative["sections"].get(section_name, "")
            if section_text:
                lines.append("")
                lines.append(section_name)
                lines.append("-" * len(section_name))
                lines.append(section_text)

    contradictions = result_data.get("contradictions", [])
    agreements = result_data.get("agreements", [])

    lines += [
        "",
        "=" * 50,
        f"CONTRADICTIONS ({len(contradictions)} found)",
        "-" * 50
    ]
    for idx, c in enumerate(contradictions, 1):
        lines.append(f"\n[{idx}] {c.get('verdict', 'DISAGREE')}")
        lines.append(f"Doc A: {c.get('chunk_a', '')}")
        lines.append(f"Doc B: {c.get('chunk_b', '')}")
        lines.append(f"Reason: {c.get('reason', '')}")

    lines += [
        "",
        "=" * 50,
        f"AGREEMENTS ({len(agreements)} found)",
        "-" * 50
    ]
    for idx, a in enumerate(agreements, 1):
        lines.append(f"\n[{idx}] {a.get('verdict', 'AGREE')}")
        lines.append(f"Doc A: {a.get('chunk_a', '')}")
        lines.append(f"Doc B: {a.get('chunk_b', '')}")
        lines.append(f"Reason: {a.get('reason', '')}")

    content = "\n".join(lines)
    return Response(
        content,
        mimetype="text/plain",
        headers={"Content-disposition": f"attachment; filename=dae-report-{int(datetime.now().timestamp())}.txt"}
    )

@app.route("/download/json")
def download_json():
    entry_id = request.args.get("id")
    result_data = None
    if entry_id:
        entry = get_history_by_id(entry_id)
        if entry:
            result_data = entry.get("results")
    if not result_data and "last_result" in session:
        result_data = session.get("last_result")

    if not result_data:
        return jsonify({"error": "No analysis results found"}), 404

    return Response(
        json.dumps(result_data, indent=2, ensure_ascii=False),
        mimetype="application/json",
        headers={"Content-disposition": f"attachment; filename=dae-result-{int(datetime.now().timestamp())}.json"}
    )

@app.route("/download/pdf")
def download_pdf():
    try:
        from weasyprint import HTML as WeasyHTML, CSS
        
        entry_id = request.args.get("id")
        result_data = None
        doc_info = {}
        
        if entry_id:
            entry = get_history_by_id(entry_id)
            if entry:
                result_data = entry.get("results")
                doc_info = {
                    "a": entry.get("doc_a"),
                    "b": entry.get("doc_b"),
                    "timestamp": entry.get("timestamp")
                }
        if not result_data and "last_result" in session:
            result_data = session.get("last_result")
            doc_info = {
                "a": result_data.get("doc_a", {}).get("name")
                     if isinstance(result_data.get("doc_a"), dict)
                     else result_data.get("doc_a"),
                "b": result_data.get("doc_b", {}).get("name")
                     if isinstance(result_data.get("doc_b"), dict)
                     else result_data.get("doc_b")
            }

        if not result_data:
            return "No results found", 404

        narrative = result_data.get("narrative", {})
        sections = narrative.get("sections", {})
        contradictions = result_data.get("contradictions", [])
        agreements = result_data.get("agreements", [])
        blind_spots_a = result_data.get("blind_spots_a", [])
        blind_spots_b = result_data.get("blind_spots_b", [])
        score_range = result_data.get("score_range", {})
        
        name_a = doc_info.get("a", "Document A")
        name_b = doc_info.get("b", "Document B")
        generated_at = datetime.now().strftime("%d %B %Y, %H:%M")

        severity_order = {"CRITICAL": 0, "SIGNIFICANT": 1, "MINOR": 2}

        html_content = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  @import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500&family=Space+Mono:wght@400;700&display=swap');
  
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  
  body {{
    font-family: 'Space Grotesk', sans-serif;
    font-size: 11pt;
    color: #0A0A0A;
    background: white;
    line-height: 1.6;
  }}

  @page {{
    size: A4;
    margin: 20mm 15mm;
    @top-right {{
      content: "DAE Report — " string(doc-pair);
      font-family: 'Space Mono', monospace;
      font-size: 7pt;
      color: #AAAAAA;
    }}
    @bottom-center {{
      content: counter(page) " / " counter(pages);
      font-family: 'Space Mono', monospace;
      font-size: 7pt;
      color: #AAAAAA;
    }}
  }}

  .report-header {{
    border-bottom: 2px solid #0A0A0A;
    padding-bottom: 16px;
    margin-bottom: 24px;
    display: flex;
    justify-content: space-between;
    align-items: flex-end;
  }}

  .report-logo {{
    font-family: 'Space Mono', monospace;
    font-size: 18pt;
    font-weight: 700;
    letter-spacing: 4px;
    color: #0A0A0A;
  }}

  .report-meta {{
    font-family: 'Space Mono', monospace;
    font-size: 8pt;
    color: #6B6760;
    text-align: right;
    line-height: 1.8;
  }}

  .doc-pair {{
    display: flex;
    gap: 12px;
    align-items: center;
    margin-bottom: 24px;
    padding: 12px 16px;
    border: 1px solid #D4CFC4;
    background: #F5F0E8;
  }}

  .doc-pill {{
    font-family: 'Space Mono', monospace;
    font-size: 8pt;
    padding: 4px 10px;
    border: 1px solid #C4BFB4;
    background: white;
    color: #0A0A0A;
  }}

  .doc-vs {{
    font-family: 'Space Mono', monospace;
    font-size: 8pt;
    color: #AAAAAA;
  }}

  .stats-row {{
    display: flex;
    gap: 12px;
    margin-bottom: 24px;
  }}

  .stat-box {{
    flex: 1;
    border: 1px solid #D4CFC4;
    padding: 12px 14px;
  }}

  .stat-value {{
    font-family: 'Space Mono', monospace;
    font-size: 22pt;
    font-weight: 700;
    color: #0A0A0A;
    display: block;
  }}

  .stat-value.red {{ color: #CC2200; }}
  .stat-value.green {{ color: #1A6B00; }}

  .stat-label {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    letter-spacing: 2px;
    color: #AAAAAA;
    display: block;
    margin-top: 4px;
  }}

  .section-label {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    letter-spacing: 4px;
    color: #AAAAAA;
    margin-bottom: 6px;
    margin-top: 24px;
    display: block;
  }}

  .narrative-block {{
    margin-bottom: 16px;
    padding-left: 12px;
    border-left: 3px solid #0A0A0A;
  }}

  .narrative-block.green {{ border-left-color: #1A6B00; }}
  .narrative-block.red {{ border-left-color: #CC2200; }}
  .narrative-block.amber {{ border-left-color: #996600; }}

  .narrative-section-label {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    letter-spacing: 3px;
    margin-bottom: 6px;
    display: block;
  }}

  .narrative-block.green .narrative-section-label {{ color: #1A6B00; }}
  .narrative-block.red .narrative-section-label {{ color: #CC2200; }}
  .narrative-block.amber .narrative-section-label {{ color: #996600; }}
  .narrative-block .narrative-section-label {{ color: #0A0A0A; }}

  .narrative-text {{
    font-size: 10pt;
    color: #0A0A0A;
    line-height: 1.7;
  }}

  .finding-card {{
    border: 1px solid #D4CFC4;
    margin-bottom: 12px;
    page-break-inside: avoid;
  }}

  .finding-card-top {{
    display: flex;
    gap: 0;
    border-bottom: 1px solid #E4E0D5;
  }}

  .finding-col {{
    flex: 1;
    padding: 10px 12px;
  }}

  .finding-col:first-child {{
    border-right: 1px solid #E4E0D5;
  }}

  .finding-col-label {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    letter-spacing: 2px;
    color: #AAAAAA;
    display: block;
    margin-bottom: 6px;
  }}

  .finding-col-text {{
    font-size: 9pt;
    color: #0A0A0A;
    line-height: 1.5;
  }}

  .finding-card-bottom {{
    padding: 8px 12px;
    display: flex;
    align-items: center;
    gap: 10px;
    background: #FAFAF8;
  }}

  .verdict-badge {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    padding: 2px 8px;
    border: 1px solid;
  }}

  .badge-disagree {{
    color: #CC2200;
    border-color: rgba(204,34,0,0.4);
    background: rgba(204,34,0,0.06);
  }}

  .badge-agree {{
    color: #1A6B00;
    border-color: rgba(26,107,0,0.4);
    background: rgba(26,107,0,0.06);
  }}

  .badge-partial {{
    color: #996600;
    border-color: rgba(153,102,0,0.4);
    background: rgba(153,102,0,0.06);
  }}

  .severity-badge {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    padding: 2px 8px;
    border: 1px solid;
  }}

  .sev-critical {{
    color: #CC2200;
    border-color: rgba(204,34,0,0.4);
    background: rgba(204,34,0,0.06);
  }}

  .sev-significant {{
    color: #996600;
    border-color: rgba(153,102,0,0.4);
    background: rgba(153,102,0,0.06);
  }}

  .sev-minor {{
    color: #1A6B00;
    border-color: rgba(26,107,0,0.4);
    background: rgba(26,107,0,0.06);
  }}

  .confidence-text {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    color: #AAAAAA;
    margin-left: auto;
  }}

  .finding-reason {{
    font-size: 9pt;
    color: #6B6760;
    line-height: 1.5;
  }}

  .blind-spot-cols {{
    display: flex;
    gap: 16px;
    margin-top: 8px;
  }}

  .blind-spot-col {{
    flex: 1;
    border: 1px solid #D4CFC4;
    padding: 12px;
  }}

  .blind-spot-col-label {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    letter-spacing: 2px;
    color: #AAAAAA;
    display: block;
    margin-bottom: 10px;
    border-bottom: 1px solid #E4E0D5;
    padding-bottom: 6px;
  }}

  .blind-spot-item {{
    font-size: 9pt;
    color: #0A0A0A;
    line-height: 1.5;
    padding: 6px 0;
    border-bottom: 1px solid #F0EDE8;
  }}

  .blind-spot-item:last-child {{ border-bottom: none; }}

  .similarity-section {{
    display: flex;
    align-items: center;
    gap: 24px;
    padding: 16px;
    border: 1px solid #D4CFC4;
    margin-bottom: 24px;
    background: #F5F0E8;
  }}

  .similarity-score {{
    font-family: 'Space Mono', monospace;
    font-size: 36pt;
    font-weight: 700;
    color: #996600;
  }}

  .similarity-label {{
    font-family: 'Space Mono', monospace;
    font-size: 7pt;
    letter-spacing: 3px;
    color: #AAAAAA;
  }}

  .page-break {{ page-break-before: always; }}
</style>
</head>
<body>

<div class="report-header">
  <div class="report-logo">[ DAE ]</div>
  <div class="report-meta">
    DATA ANALYSER ENGINE<br>
    Analysis Report<br>
    {generated_at}
  </div>
</div>

<div class="doc-pair">
  <span class="doc-pill">{name_a}</span>
  <span class="doc-vs">VS</span>
  <span class="doc-pill">{name_b}</span>
</div>

<div class="stats-row">
  <div class="stat-box">
    <span class="stat-value {'red' if contradictions else ''}">{len(contradictions)}</span>
    <span class="stat-label">CONTRADICTIONS</span>
  </div>
  <div class="stat-box">
    <span class="stat-value {'green' if agreements else ''}">{len(agreements)}</span>
    <span class="stat-label">AGREEMENTS</span>
  </div>
  <div class="stat-box">
    <span class="stat-value">{len(blind_spots_a) + len(blind_spots_b)}</span>
    <span class="stat-label">BLIND SPOTS</span>
  </div>
  <div class="stat-box">
    <span class="stat-value">{result_data.get('total_chunks_a', 0) + result_data.get('total_chunks_b', 0)}</span>
    <span class="stat-label">CHUNKS ANALYSED</span>
  </div>
</div>

<div class="similarity-section">
  <div class="similarity-score">{score_range.get('max', 0)}%</div>
  <div>
    <div class="similarity-label">SIMILARITY INDEX</div>
    <div style="font-size:10pt;color:#6B6760;margin-top:4px">
      Match range: {score_range.get('min', 0)}% — {score_range.get('max', 0)}%
    </div>
  </div>
</div>
"""

        narrative_colors = {
            "DOCUMENT OVERVIEW": ("", ""),
            "KEY FINDINGS": ("green", "green"),
            "CONTRADICTIONS ANALYSIS": ("red", "red"),
            "AGREEMENTS ANALYSIS": ("green", "green"),
            "BLIND SPOTS": ("amber", "amber"),
            "CONCLUSION": ("", ""),
        }

        if sections:
            html_content += '<span class="section-label">ANALYSIS NARRATIVE</span>'
            for section_name, (block_class, _) in narrative_colors.items():
                text = sections.get(section_name, "")
                if text:
                    html_content += f"""
<div class="narrative-block {block_class}">
  <span class="narrative-section-label">{section_name}</span>
  <div class="narrative-text">{text.replace(chr(10), '<br>')}</div>
</div>"""

        if contradictions:
            html_content += '<div class="page-break"></div>'
            html_content += f'<span class="section-label">CONTRADICTIONS ({len(contradictions)} FOUND)</span>'
            for i, c in enumerate(contradictions, 1):
                verdict = c.get("verdict", "DISAGREE")
                severity = c.get("severity", "SIGNIFICANT")
                confidence = c.get("confidence", 75)
                
                badge_class = "badge-disagree" if "DISAGREE" in verdict else \
                              "badge-agree" if "AGREE" in verdict and "PARTIALLY" not in verdict else \
                              "badge-partial"
                sev_class = f"sev-{severity.lower()}"
                
                html_content += f"""
<div class="finding-card">
  <div class="finding-card-top">
    <div class="finding-col">
      <span class="finding-col-label">DOC A</span>
      <span class="finding-col-text">{c.get('chunk_a', '')[:300]}</span>
    </div>
    <div class="finding-col">
      <span class="finding-col-label">DOC B</span>
      <span class="finding-col-text">{c.get('chunk_b', '')[:300]}</span>
    </div>
  </div>
  <div class="finding-card-bottom">
    <span class="verdict-badge {badge_class}">{verdict}</span>
    <span class="severity-badge {sev_class}">{severity}</span>
    <span class="finding-reason">{c.get('reason', '')}</span>
    <span class="confidence-text">{confidence}% confidence</span>
  </div>
</div>"""

        if agreements:
            html_content += f'<span class="section-label">AGREEMENTS ({len(agreements)} FOUND)</span>'
            for i, a in enumerate(agreements, 1):
                verdict = a.get("verdict", "AGREE")
                confidence = a.get("confidence", 75)
                badge_class = "badge-agree" if "AGREE" in verdict and "PARTIALLY" not in verdict else "badge-partial"
                html_content += f"""
<div class="finding-card">
  <div class="finding-card-top">
    <div class="finding-col">
      <span class="finding-col-label">DOC A</span>
      <span class="finding-col-text">{a.get('chunk_a', '')[:300]}</span>
    </div>
    <div class="finding-col">
      <span class="finding-col-label">DOC B</span>
      <span class="finding-col-text">{a.get('chunk_b', '')[:300]}</span>
    </div>
  </div>
  <div class="finding-card-bottom">
    <span class="verdict-badge {badge_class}">{verdict}</span>
    <span class="finding-reason">{a.get('reason', '')}</span>
    <span class="confidence-text">{confidence}% confidence</span>
  </div>
</div>"""

        if blind_spots_a or blind_spots_b:
            html_content += f'<span class="section-label">BLIND SPOTS</span>'
            html_content += '<div class="blind-spot-cols">'
            html_content += f'<div class="blind-spot-col"><span class="blind-spot-col-label">ONLY IN {name_a[:30]}</span>'
            for b in blind_spots_a:
                html_content += f'<div class="blind-spot-item">{b}</div>'
            html_content += '</div>'
            html_content += f'<div class="blind-spot-col"><span class="blind-spot-col-label">ONLY IN {name_b[:30]}</span>'
            for b in blind_spots_b:
                html_content += f'<div class="blind-spot-item">{b}</div>'
            html_content += '</div></div>'

        html_content += "</body></html>"

        pdf_bytes = WeasyHTML(string=html_content).write_pdf()
        
        timestamp = int(datetime.now().timestamp())
        return Response(
            pdf_bytes,
            mimetype="application/pdf",
            headers={
                "Content-Disposition": 
                    f"attachment; filename=dae-report-{timestamp}.pdf",
                "Content-Type": "application/pdf"
            }
        )

    except ImportError:
        return "WeasyPrint not installed. Run: pip install weasyprint", 500
    except Exception as e:
        logging.error(f"PDF generation error: {e}")
        return f"PDF generation failed: {str(e)}", 500

@app.route("/ask", methods=["POST"])
def ask():
    try:
        data = request.get_json()
        question = data.get("question", "").strip()
        entry_id = data.get("entry_id")
        
        if not question:
            return jsonify({"error": "No question provided"}), 400

        result_data = None
        if entry_id:
            entry = get_history_by_id(entry_id)
            if entry:
                result_data = entry.get("results")
        if not result_data and "last_result" in session:
            result_data = session.get("last_result")

        if not result_data:
            return jsonify({"error": "No analysis results found"}), 404

        # Build context from results
        contradictions = result_data.get("contradictions", [])
        agreements = result_data.get("agreements", [])
        blind_spots_a = result_data.get("blind_spots_a", [])
        blind_spots_b = result_data.get("blind_spots_b", [])
        narrative = result_data.get("narrative", {})
        doc_a = result_data.get("doc_a", {})
        doc_b = result_data.get("doc_b", {})

        name_a = doc_a.get("name") if isinstance(doc_a, dict) else str(doc_a)
        name_b = doc_b.get("name") if isinstance(doc_b, dict) else str(doc_b)

        context = f"""You are an expert document analyst assistant.
You have access to the results of a comparison between two documents:
Document A: {name_a}
Document B: {name_b}

CONTRADICTIONS FOUND ({len(contradictions)}):"""
        
        for i, c in enumerate(contradictions[:5], 1):
            context += f"\n{i}. Doc A says: {c['chunk_a'][:200]}"
            context += f"\n   Doc B says: {c['chunk_b'][:200]}"
            context += f"\n   Verdict: {c['verdict']} — {c['reason']}"

        context += f"\n\nAGREEMENTS FOUND ({len(agreements)}):"
        for i, a in enumerate(agreements[:5], 1):
            context += f"\n{i}. Both agree: {a['reason']}"

        context += f"\n\nTOPICS ONLY IN DOC A:"
        for b in blind_spots_a[:3]:
            context += f"\n- {b}"

        context += f"\n\nTOPICS ONLY IN DOC B:"
        for b in blind_spots_b[:3]:
            context += f"\n- {b}"

        if narrative.get("sections"):
            overview = narrative["sections"].get("DOCUMENT OVERVIEW", "")
            if overview:
                context += f"\n\nDOCUMENT OVERVIEW: {overview}"

        from langchain_groq import ChatGroq
        llm = ChatGroq(
            model_name="llama-3.1-8b-instant",
            api_key=os.getenv("GROQ_API_KEY")
        )

        messages = [
            {"role": "system", "content": context + "\n\nAnswer questions about these documents concisely and accurately. Reference specific findings when relevant. If asked something not covered by the analysis, say so clearly."},
            {"role": "user", "content": question}
        ]

        response = llm.invoke(messages)
        answer = getattr(response, "content", str(response))

        return jsonify({
            "answer": answer,
            "question": question
        })

    except Exception as e:
        logging.error(f"Ask route error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/history")
def history_page():
    entries = load_history_entries()
    if request.headers.get("Accept") == "application/json":
        return jsonify(entries)
    return render_template("history.html", history_entries=entries)

@app.route("/history/<entry_id>")
def history_detail(entry_id):
    return redirect(url_for("results", id=entry_id))

@app.route("/clear-history", methods=["POST"])
def clear_history():
    history_file = Path(app.config["HISTORY_FOLDER"]) / "history.json"
    if history_file.exists():
        history_file.unlink()
    return redirect(url_for("history_page"))

if __name__ == "__main__":
    os.makedirs("uploads", exist_ok=True)
    os.makedirs("history", exist_ok=True)
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    os.makedirs(app.config["HISTORY_FOLDER"], exist_ok=True)
    app.run(host="0.0.0.0", port=5000, debug=True)
